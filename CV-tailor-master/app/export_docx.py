from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
NEAR_BLACK = RGBColor(0x11, 0x18, 0x27)
DARK_GRAY = RGBColor(0x4B, 0x55, 0x63)


def _spacing(p, before: float = 0, after: float = 0) -> None:
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    _spacing(p, before=0, after=3)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text.upper())
    run = p.runs[0]
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = VIOLET
    _spacing(p, before=10, after=2)
    _add_horizontal_rule(doc)


def _right_tab_stop(paragraph) -> None:
    """Right-aligned tab at the full text width (6.5 inches = 9360 twips)."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")
    tabs.append(tab)
    pPr.append(tabs)


def build_tailored_docx(run_id: int, output: dict) -> BytesIO:
    """
    Build a complete send-ready CV DOCX from the tailored_cv output.
    Falls back to legacy format (summary + bullets only) if tailored_cv is absent.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    tailored_cv = output.get("tailored_cv")
    if tailored_cv:
        _build_full_cv(doc, tailored_cv)
    else:
        _build_legacy(doc, output)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ── Legacy fallback (old runs without full CV structure) ──────────────────────

def _build_legacy(doc: Document, output: dict) -> None:
    summary = output.get("tailored_summary", "").strip()
    if summary:
        h = doc.add_paragraph("PROFESSIONAL SUMMARY")
        h.runs[0].font.name = "Calibri"
        h.runs[0].font.size = Pt(9)
        h.runs[0].font.bold = True
        h.runs[0].font.color.rgb = VIOLET
        _spacing(h, after=4)
        _add_horizontal_rule(doc)
        p = doc.add_paragraph(summary)
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = NEAR_BLACK
        _spacing(p, after=14)

    bullets = output.get("tailored_bullets", []) or []
    if bullets:
        h = doc.add_paragraph("KEY EXPERIENCE")
        h.runs[0].font.name = "Calibri"
        h.runs[0].font.size = Pt(9)
        h.runs[0].font.bold = True
        h.runs[0].font.color.rgb = VIOLET
        _spacing(h, after=4)
        _add_horizontal_rule(doc)
        for bullet_text in bullets:
            p = doc.add_paragraph(str(bullet_text).strip(), style="List Bullet")
            p.runs[0].font.name = "Calibri"
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = NEAR_BLACK
            _spacing(p, after=3)


# ── Shared bullet renderer ────────────────────────────────────────────────────

def _render_bullets(doc: Document, bullets: list) -> None:
    for bullet in bullets:
        text = (
            bullet.get("text", "").strip()
            if isinstance(bullet, dict)
            else str(bullet).strip()
        )
        if text:
            bp = doc.add_paragraph(text, style="List Bullet")
            bp.runs[0].font.name = "Calibri"
            bp.runs[0].font.size = Pt(11)
            bp.runs[0].font.color.rgb = NEAR_BLACK
            _spacing(bp, after=2)


# ── Full CV builder ───────────────────────────────────────────────────────────
# Section order: Summary → Skills → Projects → Experience → Education →
#                Certifications → Awards → Leadership → Interests

def _build_full_cv(doc: Document, cv: dict) -> None:
    # ── Header ────────────────────────────────────────────────────────────────
    header = cv.get("header", {})
    name = header.get("name", "").strip()
    if name:
        p = doc.add_paragraph(name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = "Calibri"
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = NEAR_BLACK
        _spacing(p, before=0, after=4)

    contact_parts = [
        header.get(f, "").strip()
        for f in ("email", "phone", "location", "linkedin", "github", "website")
        if header.get(f, "").strip()
    ]
    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_GRAY
        _spacing(p, before=0, after=12)

    # ── Professional Summary ──────────────────────────────────────────────────
    summary = cv.get("summary", "").strip()
    if summary:
        _section_heading(doc, "Professional Summary")
        p = doc.add_paragraph(summary)
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = NEAR_BLACK
        _spacing(p, after=10)

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = cv.get("skills", [])
    if skills:
        _section_heading(doc, "Skills")
        for cat in skills:
            items = cat.get("items", [])
            if not items:
                continue
            p = doc.add_paragraph()
            r1 = p.add_run(f"{cat.get('category', '')}: ")
            r1.font.name = "Calibri"
            r1.font.size = Pt(10)
            r1.font.bold = True
            r1.font.color.rgb = NEAR_BLACK
            r2 = p.add_run(", ".join(items))
            r2.font.name = "Calibri"
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK_GRAY
            _spacing(p, after=2)

    # ── Projects (technical content leads) ───────────────────────────────────
    projects = cv.get("projects", [])
    if projects:
        _section_heading(doc, "Projects")
        for proj in projects:
            techs = proj.get("technologies", [])
            p = doc.add_paragraph()
            rn = p.add_run(proj.get("name", ""))
            rn.font.name = "Calibri"
            rn.font.size = Pt(11)
            rn.font.bold = True
            rn.font.color.rgb = NEAR_BLACK
            if techs:
                rt = p.add_run(f"  [{', '.join(techs)}]")
                rt.font.name = "Calibri"
                rt.font.size = Pt(10)
                rt.font.italic = True
                rt.font.color.rgb = DARK_GRAY
            _spacing(p, before=6, after=1)

            desc = proj.get("description", "").strip()
            if desc:
                dp = doc.add_paragraph(desc)
                dp.runs[0].font.name = "Calibri"
                dp.runs[0].font.size = Pt(10)
                dp.runs[0].font.color.rgb = DARK_GRAY
                _spacing(dp, after=1)

            _render_bullets(doc, proj.get("bullets", []))

    # ── Experience ────────────────────────────────────────────────────────────
    experience = cv.get("experience", [])
    if experience:
        _section_heading(doc, "Experience")
        for job in experience:
            start = job.get("start_date", "")
            end = job.get("end_date", "")
            date_range = f"{start} – {end}" if (start or end) else ""

            p = doc.add_paragraph()
            _right_tab_stop(p)
            r = p.add_run(job.get("company", ""))
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = NEAR_BLACK
            if date_range:
                p.add_run("\t")
                rd = p.add_run(date_range)
                rd.font.name = "Calibri"
                rd.font.size = Pt(10)
                rd.font.color.rgb = DARK_GRAY
            _spacing(p, before=8, after=1)

            title_loc = ", ".join(
                x for x in [job.get("title", ""), job.get("location", "")] if x
            )
            if title_loc:
                p2 = doc.add_paragraph(title_loc)
                r2 = p2.runs[0]
                r2.font.name = "Calibri"
                r2.font.size = Pt(10)
                r2.font.italic = True
                r2.font.color.rgb = DARK_GRAY
                _spacing(p2, after=2)

            _render_bullets(doc, job.get("bullets", []))

    # ── Education ─────────────────────────────────────────────────────────────
    education = cv.get("education", [])
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            start = edu.get("start_date", "")
            end = edu.get("end_date", "")
            date_range = f"{start} – {end}" if (start or end) else ""

            p = doc.add_paragraph()
            _right_tab_stop(p)
            r = p.add_run(edu.get("institution", ""))
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = NEAR_BLACK
            if date_range:
                p.add_run("\t")
                rd = p.add_run(date_range)
                rd.font.name = "Calibri"
                rd.font.size = Pt(10)
                rd.font.color.rgb = DARK_GRAY
            _spacing(p, before=6, after=1)

            degree_line = " ".join(
                x for x in [edu.get("degree", ""), edu.get("field", "")] if x
            )
            if edu.get("grade"):
                degree_line += f" — {edu['grade']}"
            if degree_line:
                dp = doc.add_paragraph(degree_line)
                dp.runs[0].font.name = "Calibri"
                dp.runs[0].font.size = Pt(10)
                dp.runs[0].font.italic = True
                dp.runs[0].font.color.rgb = DARK_GRAY
                _spacing(dp, after=4)

    # ── Certifications ────────────────────────────────────────────────────────
    certifications = cv.get("certifications", [])
    if certifications:
        _section_heading(doc, "Certifications")
        for cert in certifications:
            parts = [
                x for x in [
                    cert.get("name", ""),
                    cert.get("issuer", ""),
                    cert.get("date", ""),
                ] if x
            ]
            p = doc.add_paragraph(" · ".join(parts))
            if p.runs:
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = NEAR_BLACK
            _spacing(p, after=3)

    # ── Awards ────────────────────────────────────────────────────────────────
    awards = cv.get("awards", [])
    if awards:
        _section_heading(doc, "Awards")
        for award in awards:
            parts = [
                x for x in [
                    award.get("name", ""),
                    award.get("issuer", ""),
                    award.get("date", ""),
                ] if x
            ]
            p = doc.add_paragraph(" · ".join(parts))
            if p.runs:
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = NEAR_BLACK
            _spacing(p, before=4, after=1)
            desc = award.get("description", "").strip()
            if desc:
                dp = doc.add_paragraph(desc)
                dp.runs[0].font.name = "Calibri"
                dp.runs[0].font.size = Pt(10)
                dp.runs[0].font.color.rgb = DARK_GRAY
                _spacing(dp, after=3)

    # ── Leadership & Activities ────────────────────────────────────────────────
    activities = cv.get("leadership_and_activities", [])
    if activities:
        _section_heading(doc, "Leadership & Activities")
        for activity in activities:
            p = doc.add_paragraph(str(activity).strip(), style="List Bullet")
            if p.runs:
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = NEAR_BLACK
            _spacing(p, after=2)

    # ── Interests ─────────────────────────────────────────────────────────────
    interests = cv.get("interests", [])
    if interests:
        _section_heading(doc, "Interests")
        p = doc.add_paragraph(" · ".join(interests))
        if p.runs:
            p.runs[0].font.name = "Calibri"
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = DARK_GRAY
        _spacing(p, after=4)
